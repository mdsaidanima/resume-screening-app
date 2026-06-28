# you need to install all these in your terminal
# pip install streamlit
# pip install scikit-learn
# pip install python-docx
# pip install PyPDF2


import streamlit as st
import pickle
import docx  # Extract text from Word file
import PyPDF2  # Extract text from PDF
import re
import json
import traceback
from datetime import datetime
import os
import time
from pathlib import Path
import pandas as pd

BASE_DIR = Path(__file__).resolve().parent
FEEDBACK_DIR = BASE_DIR / "feedback"
FEEDBACK_FILE = FEEDBACK_DIR / "feedback_log.json"

# Load pre-trained model and TF-IDF vectorizer with caching for better performance
@st.cache_resource
def load_model():
    with open(BASE_DIR / "clf.pkl", "rb") as f:
        svc_model = pickle.load(f)
    return svc_model

@st.cache_resource
def load_vectorizer():
    with open(BASE_DIR / "tfidf.pkl", "rb") as f:
        tfidf = pickle.load(f)
    return tfidf

@st.cache_resource
def load_encoder():
    with open(BASE_DIR / "encoder.pkl", "rb") as f:
        le = pickle.load(f)
    return le

svc_model = load_model()
tfidf = load_vectorizer()
le = load_encoder()


# Category descriptions and tips
CATEGORY_DETAILS = {
    "Data Science": {
        "description": "Analyzes and interprets data to drive business decisions",
        "key_skills": ["Python", "SQL", "Machine Learning", "Statistics", "Data Visualization"],
        "tips": "Highlight experience with data analysis, modeling, statistical analysis, and big data tools"
    },
    "Java Developer": {
        "description": "Develops and maintains Java-based applications and systems",
        "key_skills": ["Java", "Spring", "OOP", "Database Design", "RESTful APIs"],
        "tips": "Include Java frameworks (Spring, Hibernate), design patterns, and project examples"
    },
    "Web Development": {
        "description": "Creates and maintains websites and web applications",
        "key_skills": ["JavaScript", "HTML/CSS", "React/Vue/Angular", "Backend", "Responsive Design"],
        "tips": "Showcase portfolio projects, frontend frameworks, and full-stack capabilities"
    },
    "Database Administrator": {
        "description": "Manages and maintains database systems and performance",
        "key_skills": ["SQL", "Database Design", "Performance Tuning", "Backup/Recovery", "Administration"],
        "tips": "Emphasize database management, optimization, security, and backup strategies"
    },
    "DevOps": {
        "description": "Focuses on deployment, automation, and infrastructure",
        "key_skills": ["Docker", "Kubernetes", "CI/CD", "Linux", "Cloud Platforms"],
        "tips": "Highlight automation tools, containerization, infrastructure-as-code, and deployment experience"
    }
}


# Function to save user feedback for model improvement
def save_feedback(predicted_category, actual_category, confidence, feedback_text="", resume_text="", filename=""):
    FEEDBACK_DIR.mkdir(parents=True, exist_ok=True)
    
    feedback_record = {
        "timestamp": datetime.now().isoformat(),
        "predicted": predicted_category,
        "actual": actual_category,
        "confidence": confidence,
        "correct": predicted_category == actual_category,
        "notes": feedback_text,
        "resume_text": resume_text,
        "filename": filename,
    }
    
    # Load existing feedback or create new list
    if FEEDBACK_FILE.exists():
        with open(FEEDBACK_FILE, 'r', encoding='utf-8') as f:
            feedback_list = json.load(f)
    else:
        feedback_list = []
    
    feedback_list.append(feedback_record)
    
    # Save updated feedback
    with open(FEEDBACK_FILE, 'w', encoding='utf-8') as f:
        json.dump(feedback_list, f, indent=2)
    
    return True


# Function to clean resume text
def cleanResume(txt):
    special_chars = r"""!"#$%&'()*+,-./:;<=>?@[\]^_`{|}~"""
    cleanText = re.sub(r'http\S+\s', ' ', txt)
    cleanText = re.sub('RT|cc', ' ', cleanText)
    cleanText = re.sub(r'#\S+\s', ' ', cleanText)
    cleanText = re.sub(r'@\S+', '  ', cleanText)
    cleanText = re.sub(r'[%s]' % re.escape(special_chars), ' ', cleanText)
    cleanText = re.sub(r'[^\x00-\x7f]', ' ', cleanText)
    cleanText = re.sub(r'\s+', ' ', cleanText)
    return cleanText


# Function to extract text from PDF
def extract_text_from_pdf(file):
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        if len(pdf_reader.pages) == 0:
            raise ValueError("PDF file has no pages")
        text = ''
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        raise ValueError(f"Failed to read PDF: {str(e)}")


# Function to extract text from DOCX
def extract_text_from_docx(file):
    try:
        doc = docx.Document(file)
        if len(doc.paragraphs) == 0:
            raise ValueError("DOCX file has no text content")
        text = ''
        for paragraph in doc.paragraphs:
            text += paragraph.text + '\n'
        return text
    except Exception as e:
        raise ValueError(f"Failed to read DOCX: {str(e)}")


# Function to extract text from TXT with explicit encoding handling
def extract_text_from_txt(file):
    # Try using utf-8 encoding for reading the text file
    try:
        text = file.read().decode('utf-8')
    except UnicodeDecodeError:
        # In case utf-8 fails, try 'latin-1' encoding as a fallback
        text = file.read().decode('latin-1')
    return text


# Function to handle file upload and extraction
def handle_file_upload(uploaded_file):
    # Validate file size (max 10 MB)
    max_size_mb = 10
    max_size_bytes = max_size_mb * 1024 * 1024
    
    if uploaded_file.size > max_size_bytes:
        raise ValueError(f"File size exceeds {max_size_mb}MB limit. Please upload a smaller file.")
    
    if '.' not in uploaded_file.name:
        raise ValueError("File must have an extension like .pdf, .docx, or .txt.")

    file_extension = uploaded_file.name.split('.')[-1].lower()
    uploaded_file.seek(0)
    
    try:
        if file_extension == 'pdf':
            text = extract_text_from_pdf(uploaded_file)
        elif file_extension == 'docx':
            text = extract_text_from_docx(uploaded_file)
        elif file_extension == 'txt':
            text = extract_text_from_txt(uploaded_file)
        else:
            raise ValueError("Unsupported file type. Please upload a PDF, DOCX, or TXT file.")
        
        # Validate extracted text
        if not text or len(text.strip()) < 10:
            raise ValueError("Could not extract meaningful text from file. Please ensure the file contains valid content.")
        
        return text
    
    except ValueError as ve:
        raise ValueError(str(ve))
    except Exception as e:
        raise ValueError(f"Error reading {file_extension.upper()} file: {str(e)}. Please ensure the file is not corrupted.")


# Function to predict the category of a resume
def pred(input_resume):
    # Preprocess the input text (e.g., cleaning, etc.)
    cleaned_text = cleanResume(input_resume)

    # Vectorize the cleaned text using the same TF-IDF vectorizer used during training
    vectorized_text = tfidf.transform([cleaned_text])

    # Convert sparse matrix to dense
    vectorized_text = vectorized_text.toarray()

    # Prediction
    predicted_category = svc_model.predict(vectorized_text)

    # Get confidence score
    try:
        # For multi-class SVM, decision_function returns scores for each class
        decision_scores = svc_model.decision_function(vectorized_text)[0]
        
        # Get all class predictions with scores
        class_scores = {}
        for idx, score in enumerate(decision_scores):
            class_name = le.classes_[idx]
            confidence = min(100, (abs(score) / (1 + abs(score))) * 100)
            class_scores[class_name] = confidence
        
        # Sort by confidence
        sorted_predictions = sorted(class_scores.items(), key=lambda x: x[1], reverse=True)
        
        # Top prediction
        top_category = sorted_predictions[0][0]
        top_confidence = sorted_predictions[0][1]
        
        # Top 3 predictions
        top_3 = sorted_predictions[:3]
        
    except:
        # Fallback
        top_category = predicted_category[0]
        top_confidence = 85.0
        top_3 = [(top_category, top_confidence)]

    return top_category, top_confidence, top_3  # Return top category, confidence, and top 3


# Function to load and calculate model metrics
@st.cache_resource
def load_model_metrics():
    metrics = {
        "total_predictions": 0,
        "correct_predictions": 0,
        "accuracy": 0.0
    }
    
    if FEEDBACK_FILE.exists():
        try:
            with open(FEEDBACK_FILE, 'r', encoding='utf-8') as f:
                feedback_list = json.load(f)
                metrics["total_predictions"] = len(feedback_list)
                metrics["correct_predictions"] = sum(1 for f in feedback_list if f.get("correct", False))
                if metrics["total_predictions"] > 0:
                    metrics["accuracy"] = (metrics["correct_predictions"] / metrics["total_predictions"]) * 100
        except:
            pass
    
    return metrics


# Streamlit app layout
def main():
    st.set_page_config(page_title="Resume Category Prediction", page_icon="📄", layout="wide")

    st.title("Resume Category Prediction App")
    st.markdown("Upload a resume in PDF, TXT, or DOCX format and get the predicted job category.")
    
    # Display model metrics
    metrics = load_model_metrics()
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("📊 Predictions Made", metrics["total_predictions"])
    with col2:
        st.metric("✅ Correct Predictions", metrics["correct_predictions"])
    with col3:
        st.metric("🎯 Model Accuracy", f"{metrics['accuracy']:.1f}%")
    
    st.markdown("---")
    
    # Advanced Analytics
    with st.expander("📈 Analytics Dashboard", expanded=False):
        st.subheader("Prediction Analytics")
        
        # Load and analyze feedback data
        if FEEDBACK_FILE.exists():
            with open(FEEDBACK_FILE, 'r', encoding='utf-8') as f:
                feedback_list = json.load(f)
            
            if len(feedback_list) > 0:
                # Prepare analytics data
                predictions = [f.get('predicted') for f in feedback_list if f.get('predicted')]
                actuals = [f.get('actual') for f in feedback_list if f.get('actual')]
                confidences = [f.get('confidence', 0) for f in feedback_list]
                
                # Category distribution
                if actuals:
                    st.write("**Distribution of Actual Categories:**")
                    category_counts = pd.Series(actuals).value_counts()
                    st.bar_chart(category_counts)
                
                # Confidence distribution
                if confidences:
                    st.write("**Confidence Score Distribution:**")
                    df_confidence = pd.DataFrame({
                        "Confidence": confidences
                    })
                    st.line_chart(df_confidence)
                
                # Accuracy by category
                if predictions and actuals:
                    st.write("**Accuracy Breakdown:**")
                    category_accuracy = {}
                    for predicted_label, actual in zip(predictions, actuals):
                        if actual not in category_accuracy:
                            category_accuracy[actual] = {"correct": 0, "total": 0}
                        category_accuracy[actual]["total"] += 1
                        if predicted_label == actual:
                            category_accuracy[actual]["correct"] += 1
                    
                    accuracy_data = []
                    for cat, data in category_accuracy.items():
                        acc = (data["correct"] / data["total"]) * 100 if data["total"] > 0 else 0
                        accuracy_data.append({"Category": cat, "Accuracy (%)": acc})
                    
                    if accuracy_data:
                        df_accuracy = pd.DataFrame(accuracy_data)
                        st.bar_chart(df_accuracy.set_index("Category"))
                
                # Recent predictions
                st.write("**Recent Predictions:**")
                recent_df = pd.DataFrame(feedback_list[-10:])
                if len(recent_df) > 0:
                    display_cols = ["timestamp", "predicted", "actual", "correct", "confidence"]
                    display_cols = [col for col in display_cols if col in recent_df.columns]
                    st.dataframe(recent_df[display_cols], use_container_width=True)
        else:
            st.info("No analytics data available yet. Submit feedback to see predictions analytics.")
    
    st.sidebar.title("Resume Prediction Toolkit")
    st.sidebar.markdown("Upload a single resume or batch multiple resumes for category prediction and scoring.")
    st.sidebar.markdown("### Quick Tips")
    st.sidebar.write("- Use the Single tab for one file")
    st.sidebar.write("- Use the Batch tab for multiple resumes")
    st.sidebar.write("- Submit feedback to improve the model")
    st.sidebar.markdown("---")
    
    st.markdown("---")
    
    # Info about retraining
    with st.expander("📚 Model Info & Retraining", expanded=False):
        st.write("""
        **About the Model:**
        - This model predicts job categories based on resume content
        - It's trained on the SVC (Support Vector Classifier) algorithm
        - Uses TF-IDF vectorization to extract features from text
        
        **Improving Model Accuracy:**
        Your feedback helps us improve! When you submit corrections:
        1. Feedback is saved in `.notebooks/feedback/feedback_log.json`
        2. Run the retraining script: `python notebooks/retrain_model.py`
        3. The model will be updated with your corrections
        4. Restart the app to use the improved model
        
        **Next Steps:**
        - Collect at least 10+ feedback entries for better retraining
        - The retraining script will backup old models automatically
        - Monitor accuracy improvements in the metrics above
        """)

    tab_single, tab_batch = st.tabs(["Single Resume", "Batch Upload"])

    with tab_single:
        st.subheader("Single Resume Prediction")
        uploaded_file = st.file_uploader("Upload a Resume", type=["pdf", "docx", "txt"], key="single_upload")

        if uploaded_file is not None:
            # Extract text from the uploaded file
            try:
                single_start = time.time()
                resume_text = handle_file_upload(uploaded_file)
                st.success("Successfully extracted the text from the uploaded resume.")

                # Display extracted text (optional)
                if st.checkbox("Show extracted text", False, key="show_text_single"):
                    st.text_area("Extracted Resume Text", resume_text, height=300)

                # Make prediction
                st.subheader("Predicted Categories")
                category, confidence, top_3 = pred(resume_text)

                # Display all top 3 predictions
                st.write("**Top 3 Predictions:**")
                for idx, (pred_category, pred_confidence) in enumerate(top_3, 1):
                    col1, col2 = st.columns([2, 1])
                    with col1:
                        st.write(f"{idx}. **{pred_category}**")
                    with col2:
                        st.write(f"Confidence: {pred_confidence:.1f}%")
                    st.progress(pred_confidence / 100)
                    st.caption(f"{pred_confidence:.1f}%")
                    st.write("")

                st.markdown("---")

                # Resume Scoring for All Categories
                st.subheader("Resume Score by Category")
                st.write("How well does your resume match each job category:")

                score_data = []
                for pred_category, pred_confidence in top_3:
                    score_data.append({"Category": pred_category, "Score": pred_confidence})

                if score_data:
                    df_scores = pd.DataFrame(score_data)
                    st.bar_chart(df_scores.set_index("Category"))

                st.markdown("---")

                # Display category details for TOP prediction
                if category in CATEGORY_DETAILS:
                    st.subheader(f"About {category} Role")
                    details = CATEGORY_DETAILS[category]

                    col1, col2 = st.columns(2)
                    with col1:
                        st.write(f"**Description:** {details['description']}")
                        st.write(f"**Key Skills:**")
                        for skill in details['key_skills']:
                            st.write(f"  • {skill}")

                    with col2:
                        st.write(f"**Tips to Improve Your Resume:**")
                        st.write(details['tips'])

                # Feedback section
                st.markdown("---")
                st.subheader("Help Us Improve 🚀")
                st.write("Is the top prediction correct? Your feedback helps us improve the model!")

                col1, col2 = st.columns(2)
                with col1:
                    is_correct = st.radio("Was the prediction correct?", ["Yes", "No"], key="feedback_correct_single")
                with col2:
                    feedback_notes = st.text_input("Any additional comments? (optional)", "", key="feedback_notes_single")

                correct_category = None
                if is_correct == "No":
                    correct_category = st.selectbox(
                        "What should the category be?",
                        options=[cat for cat in CATEGORY_DETAILS.keys() if cat != category],
                        key="correct_category_single"
                    )

                if st.button("Submit Feedback", key="submit_feedback_single"):
                    actual = category if is_correct == "Yes" else correct_category
                    save_feedback(category, actual, confidence, feedback_notes, resume_text=resume_text, filename=uploaded_file.name)
                    st.success("✅ Thank you! Your feedback has been saved and will help improve our model.")

                # Export single result
                st.markdown("---")
                st.subheader("Export Result")

                export_data = {
                    "Filename": [uploaded_file.name],
                    "Top Category": [category],
                    "Confidence (%)": [f"{confidence:.1f}"],
                    "2nd Choice": [top_3[1][0] if len(top_3) > 1 else "-"],
                    "3rd Choice": [top_3[2][0] if len(top_3) > 2 else "-"]
                }

                df_export = pd.DataFrame(export_data)
                csv = df_export.to_csv(index=False)

                st.download_button(
                    label="📥 Download Prediction as CSV",
                    data=csv,
                    file_name=f"{uploaded_file.name.split('.')[0]}_prediction.csv",
                    mime="text/csv"
                )

                elapsed = time.time() - single_start
                st.info(f"Single resume processing time: {elapsed:.2f} seconds")

            except Exception as e:
                st.error(f"Error processing the file: {e}")
                st.text(traceback.format_exc())

    with tab_batch:
        st.subheader("Batch Resume Processing")
        uploaded_files = st.file_uploader("Upload Resumes", type=["pdf", "docx", "txt"], accept_multiple_files=True, key="batch_upload")

        if uploaded_files:
            batch_start = time.time()
            st.write(f"📊 Processing {len(uploaded_files)} resumes...")

            results = []
            progress_bar = st.progress(0)
            status_text = st.empty()

            for idx, uploaded_file in enumerate(uploaded_files):
                try:
                    status_text.text(f"Processing {idx+1}/{len(uploaded_files)}: {uploaded_file.name}")
                    resume_text = handle_file_upload(uploaded_file)
                    category, confidence, top_3 = pred(resume_text)

                    results.append({
                        "filename": uploaded_file.name,
                        "top_category": category,
                        "confidence": f"{confidence:.1f}%",
                        "confidence_value": confidence,
                        "2nd_choice": top_3[1][0] if len(top_3) > 1 else "-",
                        "3rd_choice": top_3[2][0] if len(top_3) > 2 else "-"
                    })
                    progress_bar.progress((idx + 1) / len(uploaded_files))
                except Exception as e:
                    results.append({
                        "filename": uploaded_file.name,
                        "top_category": "ERROR",
                        "confidence": "-",
                        "confidence_value": 0,
                        "2nd_choice": "-",
                        "3rd_choice": "-"
                    })

            status_text.empty()
            progress_bar.empty()

            elapsed = time.time() - batch_start
            st.success(f"Batch processing complete — {len(results)} resumes processed in {elapsed:.2f} seconds.")
            st.markdown("---")

            st.subheader("Batch Summary")
            batch_col1, batch_col2 = st.columns([1, 2])
            with batch_col1:
                st.metric("Resumes Processed", len(results))
                unique_categories = len(set(r["top_category"] for r in results if r["top_category"] != "ERROR"))
                st.metric("Unique Predicted Categories", unique_categories)
                if elapsed >= 0:
                    st.metric("Elapsed Time", f"{elapsed:.2f} sec")
            with batch_col2:
                category_counts = pd.Series([r["top_category"] for r in results if r["top_category"] != "ERROR"]).value_counts()
                if not category_counts.empty:
                    st.write("**Predicted Category Distribution**")
                    st.bar_chart(category_counts)
                else:
                    st.info("No successful category predictions available for charting.")

            # Per-file confidence chart
            if any(r.get("confidence_value", 0) > 0 for r in results):
                st.markdown("---")
                st.subheader("Resume Confidence Scores")
                confidence_df = pd.DataFrame([
                    {"Resume": r["filename"], "Confidence (%)": r.get("confidence_value", 0)}
                    for r in results
                ])
                confidence_df = confidence_df.set_index("Resume")
                st.bar_chart(confidence_df)
                st.caption("Each bar shows the model's confidence for the top predicted category.")

            st.markdown("---")
            st.subheader("Batch Results")
            df_results = pd.DataFrame(results)
            df_display = df_results[["filename", "top_category", "confidence", "2nd_choice", "3rd_choice"]]
            df_display = df_display.rename(columns={
                "filename": "Filename",
                "top_category": "Predicted Category",
                "confidence": "Confidence (%)",
                "2nd_choice": "2nd Choice",
                "3rd_choice": "3rd Choice"
            })
            st.dataframe(df_display, use_container_width=True)

            csv = df_display.to_csv(index=False)
            st.download_button(
                label="📥 Download Results as CSV",
                data=csv,
                file_name="resume_predictions.csv",
                mime="text/csv"
            )

if __name__ == "__main__":
    main()